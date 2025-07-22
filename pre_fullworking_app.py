import gradio as gr
from openai import OpenAI
import PyPDF2
import docx
import os
import json
from docx import Document
from datetime import datetime

USERNAME = "admin"
PASSWORD = "letmein123"

client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

def extract_text(file_path):
    ext = os.path.splitext(file_path)[-1].lower()
    if ext == ".pdf":
        with open(file_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            return "\n".join([page.extract_text() for page in reader.pages])
    elif ext == ".docx":
        doc = docx.Document(file_path)
        return "\n".join([p.text for p in doc.paragraphs])
    elif ext == ".txt":
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    return "❌ Unsupported file type."

# ✅ Resume Generator from Personal Info
def generate_resume_from_scratch(job_description, full_name, email, phone, location, education, experience, skills):
    prompt = f"""
You are a professional resume writer. Create a resume using the personal details below:

Full Name: {full_name}
Email: {email}
Phone: {phone}
Location: {location}
Education: {education}
Experience: {experience}
Skills: {skills}

Target Job Description: {job_description}

Format the resume clearly with proper sections and highlight strengths for this job.
Only return the full formatted resume text.
"""
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.7
    )
    resume_text = response.choices[0].message.content

    filename = f"Generated_Resume_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc = Document()
    for line in resume_text.splitlines():
        doc.add_paragraph(line)
    doc.save(filename)

    return (
        "📝 Grammar: 10/10", "✅ Auto-generated resume.", "", "", "",
        filename, "", resume_text, "", []
    )

# ✅ Upload Review or Auto-Generate
def get_resume_feedback_and_rewrite(file_path, job_description, full_name, email, phone, location, education, experience, skills):
    try:
        if not file_path:
            return generate_resume_from_scratch(job_description, full_name, email, phone, location, education, experience, skills)

        resume_text = extract_text(file_path)
        prompt = f"""
You are a professional resume coach. Review the resume below and return your response strictly in this JSON format:
{{
  "scores": {{"grammar": 0-10, "structure": 0-10, "job_fit": 0-10}},
  "suggestions": ["..."],
  "rewritten_summary": "...",
  "improved_bullet_point": "...",
  "missing_keywords": ["..."],
  "rewritten_resume": "Full rewritten resume text here, with formatting and sections clearly labeled.",
  "mock_interview_questions": ["..."]
}}
Resume: {resume_text}
Target Job or Description: {job_description}
"""
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.7
        )
        content = response.choices[0].message.content
        data = json.loads(content)

        scores = (
            f"📝 Grammar: {data['scores']['grammar']}/10\n"
            f"📐 Structure: {data['scores']['structure']}/10\n"
            f"🎯 Job Fit: {data['scores']['job_fit']}/10"
        )
        suggestions = "\n- " + "\n- ".join(data["suggestions"])
        summary = data["rewritten_summary"]
        bullet = data["improved_bullet_point"]
        keywords = ", ".join(data["missing_keywords"])
        rewritten_resume_text = data["rewritten_resume"]
        interview_questions = "\n- " + "\n- ".join(data["mock_interview_questions"])
        question_list = data["mock_interview_questions"]

        doc = Document()
        for line in rewritten_resume_text.splitlines():
            doc.add_paragraph(line)
        filename = f"Rewritten_Resume_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        doc.save(filename)

        return scores, suggestions, summary, bullet, keywords, filename, resume_text, rewritten_resume_text, interview_questions, question_list
    except Exception as e:
        return [f"❌ Error: {str(e)}"] * 5 + [None, "", "", "", []]

def get_audio_feedback(audio_path, selected_question):
    try:
        if not audio_path or not os.path.exists(audio_path):
            return "❌ Please record an answer first."
        with open(audio_path, "rb") as f:
            transcript = client.audio.transcriptions.create(model="whisper-1", file=f).text
        prompt = f"""
You are an interview coach. Here's a mock interview question and the candidate's spoken answer.

Question: "{selected_question}"
Answer: "{transcript}"

Give constructive feedback on content, clarity, and relevance. Be professional and concise.
"""
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}]
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"❌ Error: {e}"

def check_login(username, password):
    if username == USERNAME and password == PASSWORD:
        return gr.update(visible=False), gr.update(visible=True), ""
    else:
        return gr.update(), gr.update(visible=False), "❌ Incorrect username or password."
# ✅ Gradio App UI with AIpply Branding
with gr.Blocks(
    title="🔐 Secure Resume App",
    css="""
        #logo img { max-width: 120px; margin-top: 10px; }
        body { background-color: #f7f9fc; }
    """
) as app:
    
    # Branding row
    with gr.Row():
        with gr.Column(scale=0.2):
            gr.Image(value="Aipply_logo.png", elem_id="logo", show_label=False, show_download_button=False)
        with gr.Column():
            gr.Markdown(
                """
                # 🤖 Welcome to **AIpply**
                > Your AI-powered Resume Assistant 🚀  
                Upload your resume or generate one from scratch with personalized feedback and interview prep!
                """,
                elem_id="header"
            )

    # Login Section
    with gr.Column(visible=True) as login_section:
        gr.Markdown("### 🔐 Login to Access the Resume Tool")
        username_input = gr.Textbox(label="Username")
        password_input = gr.Textbox(label="Password", type="password")
        login_btn = gr.Button("🔓 Login")
        login_error = gr.Textbox(label="", interactive=False)

    # Main App Section
    with gr.Column(visible=False) as main_app:
        gr.Markdown("""
## 📄 Upload Resume OR ✍️ Generate from Scratch
""")


You can either:
- 📤 Upload your existing resume (`.pdf`, `.docx`, or `.txt`)  
**OR**
- 🧑‍💼 Fill in your info & 🎯 Job Target below to auto-generate a professional resume!
""")

        with gr.Row():
            resume_file = gr.File(label="📄 Upload Resume (.pdf, .docx, .txt)", type="filepath")
            job_input = gr.Textbox(label="📋 Job Description or Target Job", lines=4)

        gr.Markdown("### 🧑‍💼️ Personal Info (for resume generation)")
        full_name = gr.Textbox(label="👤 Full Name")
        email = gr.Textbox(label="📧 Email")
        phone = gr.Textbox(label="📱 Phone Number")
        location = gr.Textbox(label="🌍 Location")
        education = gr.Textbox(label="🎓 Education", lines=2)
        experience = gr.Textbox(label="💼 Experience", lines=2)
        skills = gr.Textbox(label="🛠 Skills (comma-separated)")

        submit = gr.Button("🧠 Analyze / Generate Resume")

        with gr.Row():
            scores_out = gr.Textbox(label="📊 Scores")
            suggestions_out = gr.Textbox(label="🛠 Suggestions")
        with gr.Row():
            summary_out = gr.Textbox(label="✍️ Rewritten Summary")
            bullet_out = gr.Textbox(label="🔁 Improved Bullet Point")
            keywords_out = gr.Textbox(label="❗ Missing Keywords")
        file_out = gr.File(label="⬇️ Download Resume (.docx)")
        gr.Markdown("## 🆚 Resume Comparison")
        with gr.Row():
            original_out = gr.Textbox(label="📄 Original Resume", lines=20)
            rewritten_out = gr.Textbox(label="✍️ Rewritten Resume", lines=20)
        gr.Markdown("## 🎤 Mock Interview Questions")
        interview_out = gr.Textbox(label="Interview Questions", lines=10)
        gr.Markdown("## 🎙️ Practice Interview Responses")
        with gr.Row():
            question_dropdown = gr.Dropdown(label="🧠 Select a Mock Question", choices=[], interactive=True)
            audio_input = gr.Audio(sources="microphone", type="filepath", label="🎤 Record Your Answer")
        feedback_btn = gr.Button("🔍 Get Feedback on Answer")
        feedback_out = gr.Textbox(label="🧠 GPT Feedback")

    login_btn.click(check_login, [username_input, password_input], [login_section, main_app, login_error])
    submit.click(get_resume_feedback_and_rewrite, [
        resume_file, job_input, full_name, email, phone, location, education, experience, skills
    ], [
        scores_out, suggestions_out, summary_out, bullet_out,
        keywords_out, file_out, original_out, rewritten_out,
        interview_out, question_dropdown
    ])
    feedback_btn.click(get_audio_feedback, [audio_input, question_dropdown], feedback_out)

if __name__ == "__main__":
    print("✅ Launching app...")
    app.queue().launch(
        server_name="0.0.0.0",
        server_port=int(os.environ.get("PORT", 8080))
    )
